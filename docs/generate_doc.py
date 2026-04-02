#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys, io
# Force UTF-8 output on Windows consoles (cp1251 doesn't handle all Unicode)
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')
"""
NexusCRM — Генератор проектной документации
Колледж Хекслет · Шафиев Артём · группа 13 ТИС · 2026

Запуск:
    python generate_doc.py

Результат:
    NexusCRM_Documentation.docx  — исходный редактируемый файл
    NexusCRM_Documentation.pdf   — готовый PDF (через Microsoft Word / LibreOffice)
"""

import os, sys
from pathlib import Path
from io import BytesIO

# ── Зависимости ────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.shared import Inches
except ImportError:
    sys.exit("Установите python-docx:  pip install python-docx")

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
    MPLOT = True
except ImportError:
    MPLOT = False
    print("[WARNING] matplotlib не найден — диаграммы будут пропущены")

DOCS_DIR   = Path(__file__).parent
OUT_DOCX   = DOCS_DIR / 'NexusCRM_Documentation.docx'
OUT_PDF    = DOCS_DIR / 'NexusCRM_Documentation.pdf'

FONT       = 'Times New Roman'
SZ_BODY    = 12
SZ_FOOT    = 10
SZ_TABLE   = 11
SZ_CAPTION = 12

# ══════════════════════════════════════════════════════════════════════════════
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ══════════════════════════════════════════════════════════════════════════════

def _rpr_font(element, name):
    """Принудительно устанавливает шрифт в XML rPr для кириллицы."""
    try:
        rPr = element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts.set(qn(attr), name)
    except Exception:
        pass


def _set_run(run, name=None, size=None, bold=False, italic=False):
    name = name or FONT
    size = size or SZ_BODY
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    _rpr_font(run._r, name)


def _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
          space_before=0, space_after=0, left_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment        = align
    pf.space_before     = Pt(space_before)
    pf.space_after      = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(1.25) if indent_first else Cm(0)
    if left_indent is not None:
        pf.left_indent = left_indent
    return p


def add_body(doc, text, indent_first=True):
    p = _para(doc, indent_first=indent_first)
    r = p.add_run(text)
    _set_run(r)
    return p


def add_body_run(doc, parts, indent_first=True):
    """parts = list of (text, bold, italic)"""
    p = _para(doc, indent_first=indent_first)
    for text, bold, italic in parts:
        r = p.add_run(text)
        _set_run(r, bold=bold, italic=italic)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run('')
    _set_run(r)
    return p


def add_structural(doc, text):
    """ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т.д. — ПРОПИСНЫЕ ЖИРНЫЕ ПО ЦЕНТРУ"""
    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run(text.upper())
    _set_run(r, bold=True)
    add_blank(doc)
    return p


def add_section(doc, number, text):
    """1 НАЗВАНИЕ РАЗДЕЛА — прописные жирные по центру, без отступа"""
    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run(f'{number} {text.upper()}')
    _set_run(r, bold=True)
    add_blank(doc)
    return p


def add_paragraph_heading(doc, number, text):
    """1.1 Заголовок параграфа — строчные с прописной, жирные по центру"""
    add_blank(doc)
    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run(f'{number} {text}')
    _set_run(r, bold=True)
    add_blank(doc)
    return p


def add_list_item(doc, text, prefix='–', numbered=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.left_indent       = Cm(1.25)
    if numbered:
        lead = f'{numbered}) '
    else:
        lead = f'{prefix} '
    r = p.add_run(lead + text)
    _set_run(r)
    return p


def add_table_caption(doc, number, title):
    """Таблица N – Название — по левому краю, без отступа"""
    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    r = p.add_run(f'Таблица {number} \u2013 {title}')
    _set_run(r)
    return p


def add_figure_caption(doc, number, title, note=None):
    """Рисунок N – Название — по центру, без отступа"""
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run(f'Рисунок {number} \u2013 {title}')
    _set_run(r)
    if note:
        pn = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
        rn = pn.add_run(f'Примечание \u2013 {note}')
        _set_run(rn)
    add_blank(doc)
    return p


def embed_image(doc, buf, width_cm=14.0):
    """Вставляет изображение из BytesIO, выравнивание по центру."""
    add_blank(doc)
    p = doc.add_paragraph()
    p.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before     = Pt(0)
    p.paragraph_format.space_after      = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))
    return p


def add_pagebreak(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# НАСТРОЙКА ДОКУМЕНТА
# ══════════════════════════════════════════════════════════════════════════════

def setup_document():
    doc = Document()

    # Поля: верх 2 см, лево 3 см, низ 2 см, право 1,5 см
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(1.5)
        sec.footer_distance = Cm(1.25)

    # Стиль по умолчанию
    normal = doc.styles['Normal']
    normal.font.name  = FONT
    normal.font.size  = Pt(SZ_BODY)
    normal.paragraph_format.space_before     = Pt(0)
    normal.paragraph_format.space_after      = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.JUSTIFY
    _rpr_font(normal.element, FONT)

    return doc


def add_page_numbers(doc):
    """Добавляет нумерацию страниц по центру в нижнем колонтитуле."""
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        p = footer.paragraphs[0]
        p.clear()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run.font.name = FONT
        run.font.size = Pt(SZ_FOOT)
        _rpr_font(run._r, FONT)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def hide_first_page_number(doc):
    """Убирает отображение номера на первой странице (титул)."""
    for sec in doc.sections:
        titlePg = OxmlElement('w:titlePg')
        sec._sectPr.append(titlePg)
        pfirst = OxmlElement('w:pgNumType')
        pfirst.set(qn('w:start'), '1')
        sec._sectPr.append(pfirst)


# ══════════════════════════════════════════════════════════════════════════════
# ДИАГРАММЫ (matplotlib)
# ══════════════════════════════════════════════════════════════════════════════

GRAY_BG   = '#1e1e2e'
ACCENT    = '#6366f1'
ACCENT2   = '#8b5cf6'
ACCENT3   = '#10b981'
FG        = '#e2e8f0'
BORDER    = '#374151'


def fig_to_buf(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


def make_architecture_diagram():
    """Рисунок 1 — Архитектура приложения."""
    fig, ax = plt.subplots(figsize=(12, 6))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis('off')
    ax.set_facecolor('white')

    def box(x, y, w, h, color, label, sub='', fontsize=9):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle='round,pad=0.1',
                              linewidth=1.2, edgecolor='#555',
                              facecolor=color)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.12 if sub else 0), label,
                ha='center', va='center', fontsize=fontsize,
                fontweight='bold', color='black')
        if sub:
            ax.text(x + w/2, y + h/2 - 0.2, sub,
                    ha='center', va='center', fontsize=7, color='#444')

    def arrow(x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.4))

    # Слой 1 — Хранение данных
    ax.text(0.3, 0.55, 'Слой хранения данных', fontsize=8, color='#333',
            rotation=90, va='center')
    box(0.7, 0.2, 3.2, 0.7, '#fef3c7', 'localStorage', 'nexus_crm_clients')

    # Слой 2 — Логика
    ax.text(0.3, 2.5, 'Слой бизнес-логики (JavaScript ES-модули)', fontsize=8,
            color='#333', rotation=90, va='center')
    box(0.7, 1.4, 1.3, 0.7, '#dbeafe', 'storage.js', 'Хранение данных')
    box(2.1, 1.4, 1.3, 0.7, '#dbeafe', 'utils.js',   'Утилиты')
    box(0.7, 2.3, 1.3, 0.7, '#e0e7ff', 'main.js',    'Главная')
    box(2.1, 2.3, 1.3, 0.7, '#e0e7ff', 'form.js',    'Форма')
    box(3.5, 2.3, 1.3, 0.7, '#e0e7ff', 'stats.js',   'Статистика')

    # Слой 3 — Представление
    ax.text(0.3, 4.5, 'Слой представления (HTML + CSS)', fontsize=8,
            color='#333', rotation=90, va='center')
    box(0.7,  3.8, 1.3, 0.7, '#d1fae5', 'index.html',  'Клиенты')
    box(2.1,  3.8, 1.3, 0.7, '#d1fae5', 'form.html',   'Форма')
    box(3.5,  3.8, 1.3, 0.7, '#d1fae5', 'stats.html',  'Статистика')
    box(4.9,  3.8, 1.3, 0.7, '#d1fae5', 'about.html',  'О проекте')
    box(0.7,  4.7, 5.5, 0.6, '#f0fdf4', 'css/style.css — Общая таблица стилей',
        fontsize=8)

    # Слой 4 — Пользователь
    box(0.7, 5.5, 5.5, 0.4, '#fce7f3', 'Пользователь (браузер)', fontsize=9)

    # Стрелки
    arrow(1.95, 0.9, 1.35, 1.4)    # localStorage → storage.js
    arrow(1.35, 2.1, 1.35, 2.3)    # storage.js → main.js
    arrow(2.75, 2.1, 2.75, 2.3)    # utils.js → form.js
    arrow(1.35, 3.0, 1.35, 3.8)    # main.js → index.html
    arrow(2.75, 3.0, 2.75, 3.8)    # form.js → form.html
    arrow(4.15, 3.0, 4.15, 3.8)    # stats.js → stats.html
    arrow(3.5,  5.5, 3.5,  4.5)    # Пользователь → HTML

    # Легенда
    legend_items = [
        mpatches.Patch(facecolor='#fce7f3', edgecolor='#555', label='Пользователь'),
        mpatches.Patch(facecolor='#d1fae5', edgecolor='#555', label='HTML-страницы'),
        mpatches.Patch(facecolor='#e0e7ff', edgecolor='#555', label='Прикладная логика'),
        mpatches.Patch(facecolor='#dbeafe', edgecolor='#555', label='Инфраструктура'),
        mpatches.Patch(facecolor='#fef3c7', edgecolor='#555', label='Хранилище данных'),
    ]
    ax.legend(handles=legend_items, loc='center right', fontsize=7,
              framealpha=0.9, bbox_to_anchor=(11.8, 3.0))

    ax.set_title('Архитектура приложения NexusCRM', fontsize=11, pad=10)
    return fig_to_buf(fig)


def make_navigation_diagram():
    """Рисунок 2 — Схема переходов между страницами."""
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_facecolor('white')

    def box(x, y, w, h, color, label, sub=''):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle='round,pad=0.15',
                              linewidth=1.3, edgecolor='#444',
                              facecolor=color)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.1 if sub else 0), label,
                ha='center', va='center', fontsize=10, fontweight='bold')
        if sub:
            ax.text(x + w/2, y + h/2 - 0.2, sub,
                    ha='center', va='center', fontsize=7.5, color='#555')

    def arr(x1, y1, x2, y2, label='', col='#555'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=col, lw=1.4))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx + 0.05, my, label, fontsize=7.5, color=col, ha='left')

    box(0.3,  1.8, 2.2, 1.0, '#d1fae5', 'index.html', 'Список клиентов')
    box(3.5,  1.8, 2.2, 1.0, '#e0e7ff', 'form.html',  'Добавить/Редактировать')
    box(7.0,  1.8, 2.2, 1.0, '#fef3c7', 'stats.html', 'Статистика')
    box(3.5,  3.5, 2.2, 1.0, '#fce7f3', 'about.html', 'О проекте')
    box(0.3,  0.2, 2.2, 1.0, '#f0fdf4', 'Браузер',    'Открыть файл / URL')

    arr(1.4, 1.2,  1.4, 1.8, '')
    arr(2.5, 2.3,  3.5, 2.3, 'Добавить')
    arr(3.5, 2.7,  2.5, 2.7, 'Сохранить')
    arr(2.5, 2.5,  3.5, 2.0, 'Редактировать')
    arr(6.0, 2.3,  7.0, 2.3, 'Статистика', '#888')
    arr(7.0, 2.7,  6.0, 2.7, 'Назад', '#888')
    arr(2.5, 4.0,  1.4, 3.0, 'О проекте', '#888')
    arr(4.6, 3.5,  4.6, 2.8, 'О проекте', '#888')

    ax.set_title('Схема переходов между страницами NexusCRM', fontsize=11, pad=10)
    return fig_to_buf(fig)


def make_algorithm_diagram():
    """Рисунок 3 — Алгоритм инициализации главной страницы."""
    fig, ax = plt.subplots(figsize=(8, 9))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('white')

    def rect_box(x, y, w, h, color, text, fontsize=9):
        r = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.1',
                           linewidth=1.2, edgecolor='#444', facecolor=color)
        ax.add_patch(r)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=fontsize, wrap=True,
                multialignment='center')

    def diamond(x, y, w, h, color, text):
        cx, cy = x + w/2, y + h/2
        pts = [(cx, y+h), (x+w, cy), (cx, y), (x, cy)]
        from matplotlib.patches import Polygon
        poly = Polygon(pts, closed=True, linewidth=1.2,
                       edgecolor='#444', facecolor=color)
        ax.add_patch(poly)
        ax.text(cx, cy, text, ha='center', va='center', fontsize=8)

    def arr(x1, y1, x2, y2, label=''):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#333', lw=1.3))
        if label:
            ax.text((x1+x2)/2 + 0.1, (y1+y2)/2, label, fontsize=7.5)

    # Блоки алгоритма
    rect_box(2.5, 8.3, 3.0, 0.5, '#fce7f3', 'НАЧАЛО', 10)
    arr(4.0, 8.3, 4.0, 7.8)
    rect_box(2.0, 7.3, 4.0, 0.5, '#dbeafe', 'Загрузка страницы index.html')
    arr(4.0, 7.3, 4.0, 6.8)
    rect_box(2.0, 6.3, 4.0, 0.5, '#e0e7ff', 'init(): initClients()')
    arr(4.0, 6.3, 4.0, 5.8)
    diamond(2.5, 5.2, 3.0, 0.6, '#fef3c7', 'localStorage\nпуст?')
    # Да
    arr(2.5, 5.5, 1.5, 5.5, '')
    rect_box(0.2, 5.2, 1.2, 0.6, '#d1fae5', 'Сеять\nдемо-данные', 7)
    arr(0.8, 5.2, 0.8, 4.5, '')
    arr(0.8, 4.5, 3.5, 4.5, '')
    # Нет
    arr(5.5, 5.5, 6.0, 5.5, 'Нет')
    rect_box(6.0, 5.2, 1.7, 0.6, '#d1fae5', 'Читать\nclients[]', 7)
    arr(6.85, 5.2, 6.85, 4.5, '')
    arr(6.85, 4.5, 5.5, 4.5, '')

    rect_box(2.0, 4.1, 4.0, 0.5, '#e0e7ff', 'renderClients(): фильтр + рендер карточек')
    arr(4.0, 4.1, 4.0, 3.6)
    rect_box(2.0, 3.1, 4.0, 0.5, '#e0e7ff', 'setupSearch(), setupFilters()')
    arr(4.0, 3.1, 4.0, 2.6)
    diamond(2.5, 2.0, 3.0, 0.6, '#fef3c7', 'Событие\npageshow\n(persisted)?')
    arr(4.0, 2.0, 4.0, 1.5, 'Нет → стоп')
    arr(5.5, 2.3, 6.2, 2.3, 'Да')
    rect_box(6.2, 2.0, 1.5, 0.6, '#fce7f3', 'reinit()\n(без событий)', 7)
    arr(4.0, 1.5, 4.0, 1.0)
    rect_box(2.5, 0.5, 3.0, 0.5, '#fce7f3', 'КОНЕЦ', 10)

    ax.set_title('Алгоритм инициализации главной страницы', fontsize=11, pad=10)
    return fig_to_buf(fig)


# ══════════════════════════════════════════════════════════════════════════════
# ТАБЛИЦЫ
# ══════════════════════════════════════════════════════════════════════════════

def _tbl_set_style(table):
    """Применяет базовый стиль к таблице."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)


def _cell_text(cell, text, bold=False, center=False, size=SZ_TABLE, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                    if center else WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(str(text))
    r.font.name   = FONT
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    _rpr_font(r._r, FONT)


def make_table1(doc):
    """Таблица 1 — Сравнительный анализ CRM-решений"""
    add_table_caption(doc, 1,
        'Сравнительный анализ существующих CRM-решений')
    headers = ['Критерий', 'Bitrix24', 'AmoCRM', 'Salesforce', 'NexusCRM']
    rows = [
        ['Тип размещения',     'Облачный SaaS', 'Облачный SaaS', 'Облачный SaaS', 'Клиентский (браузер)'],
        ['Серверная часть',    'Обязательна',   'Обязательна',   'Обязательна',   'Не требуется'],
        ['Интернет для работы','Обязателен',     'Обязателен',    'Обязателен',    'Не требуется'],
        ['Бесплатный план',    'Ограниченный',  'Отсутствует',   'Отсутствует',   'Полный'],
        ['Сложность освоения', 'Высокая',       'Средняя',       'Высокая',       'Низкая'],
        ['Кастомизация',       'Высокая',       'Средняя',       'Высокая',       'Ограниченная'],
        ['Экспорт данных',     'Есть',          'Есть',          'Есть',          'Не реализован'],
    ]
    col_widths = [Cm(4.0), Cm(2.7), Cm(2.7), Cm(2.7), Cm(3.5)]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Заголовок
    for j, h in enumerate(headers):
        _cell_text(tbl.rows[0].cells[j], h, bold=True, center=True, size=SZ_TABLE)

    # Нумерация столбцов
    num_row = tbl.add_row()
    for j in range(len(headers)):
        _cell_text(num_row.cells[j], str(j+1), center=True,
                   size=SZ_TABLE, italic=True)

    # Данные
    for i, row_data in enumerate(rows):
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            _cell_text(row.cells[j], val, center=(j > 0), size=SZ_TABLE)

    # Ширина столбцов
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[j].twips)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # Примечание
    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    r = p.add_run('Примечание \u2013 составлено автором на основании проведённого исследования')
    _set_run(r)
    add_blank(doc)
    return tbl


def make_table2(doc):
    """Таблица 2 — Поля объекта клиента"""
    add_table_caption(doc, 2, 'Поля объекта клиента в хранилище данных')
    headers = ['Поле', 'Тип данных', 'Обязательное', 'Описание']
    rows = [
        ['id',        'string', 'Да',  'Уникальный идентификатор: метка времени + случайное число'],
        ['fullName',  'string', 'Да',  'Полное имя клиента'],
        ['phone',     'string', 'Да',  'Номер телефона'],
        ['email',     'string', 'Да',  'Адрес электронной почты (валидируется по формату)'],
        ['status',    'string', 'Да',  'Статус сделки: new, active или completed'],
        ['comment',   'string', 'Нет', 'Текстовый комментарий, до 500 символов'],
        ['createdAt', 'string', 'Да',  'Дата создания в формате ISO 8601'],
    ]
    col_widths = [Cm(2.6), Cm(2.4), Cm(2.8), Cm(7.6)]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'

    for j, h in enumerate(headers):
        _cell_text(tbl.rows[0].cells[j], h, bold=True, center=True, size=SZ_TABLE)

    num_row = tbl.add_row()
    for j in range(len(headers)):
        _cell_text(num_row.cells[j], str(j+1), center=True, size=SZ_TABLE, italic=True)

    for row_data in rows:
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            _cell_text(row.cells[j], val, center=(j in (1, 2)), size=SZ_TABLE)

    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[j].twips)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    r = p.add_run('Примечание \u2013 составлено автором на основании проведённого исследования')
    _set_run(r)
    add_blank(doc)
    return tbl


def make_table3(doc):
    """Таблица 3 — Файловая структура проекта"""
    add_table_caption(doc, 3, 'Файловая структура проекта NexusCRM')
    headers = ['Файл / директория', 'Назначение']
    rows = [
        ['index.html',      'Главная страница: отображение списка клиентов'],
        ['form.html',       'Форма добавления и редактирования клиента'],
        ['stats.html',      'Страница статистики и аналитики по базе клиентов'],
        ['about.html',      'Информационная страница о проекте'],
        ['css/style.css',   'Основная таблица стилей: тёмная тема, переменные, адаптивность'],
        ['js/main.js',      'Логика главной страницы: рендер, поиск, фильтр, CRUD'],
        ['js/storage.js',   'Слой хранения данных: операции с localStorage'],
        ['js/form.js',      'Логика формы: валидация, добавление, редактирование'],
        ['js/stats.js',     'Расчёт и отображение статистических показателей'],
        ['js/utils.js',     'Общие вспомогательные функции для всех модулей'],
    ]
    col_widths = [Cm(4.5), Cm(11.0)]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'

    for j, h in enumerate(headers):
        _cell_text(tbl.rows[0].cells[j], h, bold=True, center=True, size=SZ_TABLE)

    num_row = tbl.add_row()
    for j in range(2):
        _cell_text(num_row.cells[j], str(j+1), center=True, size=SZ_TABLE, italic=True)

    for row_data in rows:
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            _cell_text(row.cells[j], val, size=SZ_TABLE)

    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[j].twips)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    r = p.add_run('Примечание \u2013 составлено автором на основании проведённого исследования')
    _set_run(r)
    add_blank(doc)
    return tbl


def make_table4(doc):
    """Таблица 4 — Стек технологий"""
    add_table_caption(doc, 4, 'Стек технологий проекта NexusCRM')
    headers = ['Технология', 'Версия', 'Назначение']
    rows = [
        ['HTML5',               '\u2014', 'Структура страниц, семантическая разметка'],
        ['CSS3',                '\u2014', 'Стилизация, тёмная цветовая схема, адаптивная вёрстка'],
        ['JavaScript (ES2020+)', '\u2014', 'Бизнес-логика, ES-модули (import/export)'],
        ['Web Storage API',     '\u2014', 'Персистентное хранение данных в localStorage'],
        ['Google Fonts API',    '\u2014', 'Подключение шрифта Inter для интерфейса'],
    ]
    col_widths = [Cm(4.0), Cm(2.5), Cm(9.0)]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'

    for j, h in enumerate(headers):
        _cell_text(tbl.rows[0].cells[j], h, bold=True, center=True, size=SZ_TABLE)

    num_row = tbl.add_row()
    for j in range(3):
        _cell_text(num_row.cells[j], str(j+1), center=True, size=SZ_TABLE, italic=True)

    for row_data in rows:
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            _cell_text(row.cells[j], val, center=(j==1), size=SZ_TABLE)

    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[j].twips)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    add_blank(doc)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    r = p.add_run('Примечание \u2013 составлено автором на основании проведённого исследования')
    _set_run(r)
    add_blank(doc)
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# СОДЕРЖАНИЕ (оглавление)
# ══════════════════════════════════════════════════════════════════════════════

def add_contents_page(doc):
    add_structural(doc, 'СОДЕРЖАНИЕ')

    toc_entries = [
        ('ВВЕДЕНИЕ', '3'),
        ('1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ', '4'),
        ('   1.1 Актуальность разработки', '4'),
        ('   1.2 Анализ существующих CRM-решений', '5'),
        ('   1.3 Цель и задачи проекта', '6'),
        ('2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ', '7'),
        ('   2.1 Концептуальная архитектура приложения', '7'),
        ('   2.2 Модель данных клиента', '9'),
        ('   2.3 Структура файлов проекта', '10'),
        ('3 РЕАЛИЗАЦИЯ', '11'),
        ('   3.1 Стек технологий', '11'),
        ('   3.2 Описание программных модулей', '12'),
        ('   3.3 Пользовательский интерфейс', '14'),
        ('   3.4 Алгоритм работы приложения', '16'),
        ('4 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ', '18'),
        ('   4.1 Требования к запуску и установке', '18'),
        ('   4.2 Основные сценарии использования', '19'),
        ('   4.3 Ограничения и перспективы развития', '21'),
        ('ЗАКЛЮЧЕНИЕ', '22'),
        ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '23'),
        ('Приложение 1', '25'),
        ('Приложение 2', '26'),
    ]

    for title, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before     = Pt(0)
        p.paragraph_format.space_after      = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)

        r1 = p.add_run(title)
        _set_run(r1)

        # Точки-заполнители и номер страницы справа
        tab = p.add_run('\t' + page)
        _set_run(tab)

        # Вкладка с заполнением
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_stop = OxmlElement('w:tab')
        tab_stop.set(qn('w:val'), 'right')
        tab_stop.set(qn('w:leader'), 'dot')
        tab_stop.set(qn('w:pos'), '9072')  # 16 см в twips
        tabs.append(tab_stop)
        pPr.append(tabs)

    add_pagebreak(doc)


# ══════════════════════════════════════════════════════════════════════════════
# ТИТУЛЬНЫЙ ЛИСТ
# ══════════════════════════════════════════════════════════════════════════════

def add_title_page(doc):
    def centered(text, bold=False, size=SZ_BODY):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
        r = p.add_run(text)
        _set_run(r, bold=bold, size=size)
        return p

    centered('ТОО «Колледж Хекслет»', bold=False)
    add_blank(doc)
    centered('Направление подготовки: Информационные системы', bold=False)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    centered('ПРОЕКТНАЯ РАБОТА', bold=True, size=14)
    add_blank(doc)
    add_blank(doc)

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run('Тема: «Разработка веб-приложения NexusCRM\n'
                  'для управления базой клиентов малого бизнеса»')
    _set_run(r, bold=False, size=SZ_BODY)

    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)

    # Блок справа
    def right_line(text):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, indent_first=False)
        r = p.add_run(text)
        _set_run(r)

    right_line('Выполнил: студент 2 курса,')
    right_line('группа 13 ТИС')
    right_line('Шафиев Артём')
    add_blank(doc)
    right_line('Руководитель:')
    right_line('преподаватель Колледжа Хекслет')

    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)

    centered('Алматы 2026')
    add_pagebreak(doc)


# ══════════════════════════════════════════════════════════════════════════════
# ОСНОВНОЙ ТЕКСТ ДОКУМЕНТА
# ══════════════════════════════════════════════════════════════════════════════

def add_introduction(doc):
    add_structural(doc, 'ВВЕДЕНИЕ')

    add_body(doc,
        'В современных условиях развития малого бизнеса одним из ключевых '
        'инструментов повышения эффективности работы с клиентами является система '
        'управления взаимоотношениями с клиентами (CRM, от англ. Customer Relationship '
        'Management). CRM-системы позволяют структурировать информацию о клиентах, '
        'отслеживать статусы сделок, вести историю взаимодействий и анализировать '
        'ключевые показатели работы команды продаж [1].')

    add_body(doc,
        'Вместе с тем существующие на рынке CRM-решения — Bitrix24, AmoCRM, Salesforce '
        '— обладают рядом существенных недостатков для субъектов малого бизнеса: высокая '
        'стоимость лицензии или подписки, избыточный функционал, обязательное наличие '
        'серверной инфраструктуры и постоянного интернет-соединения, а также сложность '
        'первоначальной настройки. Нередко предприниматели и небольшие команды ведут '
        'клиентскую базу в таблицах Microsoft Excel или Google Sheets, что не обеспечивает '
        'должного уровня удобства, статусной модели и аналитических инструментов [2].')

    add_body(doc,
        'Данная проектная работа посвящена разработке клиентского веб-приложения '
        'NexusCRM \u2014 системы управления базой клиентов для малого бизнеса, '
        'реализованной на нативных технологиях HTML5, CSS3 и JavaScript с применением '
        'ES-модулей. Приложение функционирует полностью в браузере без необходимости '
        'развёртывания серверной части и подключения к интернету при работе с данными.')

    add_body(doc,
        'Целью проекта является разработка функционального веб-приложения для ведения '
        'базы клиентов с реализацией операций добавления, редактирования, удаления, '
        'поиска и фильтрации записей, а также модуля статистической аналитики.')

    add_body(doc,
        'Для достижения поставленной цели были определены следующие задачи:')
    add_list_item(doc, 'провести анализ предметной области и существующих CRM-решений;')
    add_list_item(doc, 'спроектировать архитектуру клиентского веб-приложения;')
    add_list_item(doc, 'реализовать слой хранения данных на основе localStorage браузера;')
    add_list_item(doc, 'разработать пользовательский интерфейс с тёмной цветовой схемой '
                       'и адаптивной вёрсткой;')
    add_list_item(doc, 'реализовать полный набор операций управления клиентскими данными;')
    add_list_item(doc, 'разработать модуль статистики и аналитики по клиентской базе;')
    add_list_item(doc, 'обеспечить защиту от межсайтового скриптинга (XSS) и '
                       'устойчивость к ошибкам хранилища данных.')

    add_body(doc,
        'Объектом исследования являются процессы управления клиентскими данными '
        'в условиях малого бизнеса.')

    add_body(doc,
        'Предметом исследования служат технологии разработки клиентских веб-приложений '
        'для автоматизации CRM-процессов.')

    add_body(doc,
        'Практическая значимость работы определяется готовностью разработанного '
        'приложения к применению в реальных условиях: оно не требует установки '
        'дополнительного программного обеспечения, функционирует без подключения к '
        'интернету и может быть развёрнуто на любом статическом хостинге.')

    add_pagebreak(doc)


def add_section1(doc):
    add_section(doc, '1', 'АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ')

    # 1.1
    add_paragraph_heading(doc, '1.1', 'Актуальность разработки')

    add_body(doc,
        'В условиях цифровизации экономики управление взаимоотношениями с клиентами '
        'становится критически важным элементом успешной деятельности предприятий любого '
        'масштаба. По данным аналитических агентств, использование CRM-систем повышает '
        'производительность менеджеров по продажам на 26\u201334%, увеличивает конверсию '
        'лидов в сделки и обеспечивает систематизацию клиентской базы [3].')

    add_body(doc,
        'Субъекты малого бизнеса \u2014 индивидуальные предприниматели, небольшие '
        'команды продаж, фриланс-студии и агентства \u2014 нередко сталкиваются с '
        'проблемой отсутствия доступного инструмента для ведения клиентской базы. '
        'Корпоративные CRM-решения предъявляют высокие требования к инфраструктуре, '
        'требуют постоянного интернет-соединения и обладают избыточным функционалом, '
        'не востребованным малым бизнесом.')

    add_body(doc,
        'Специфика работы небольших команд предполагает необходимость простого и '
        'быстрого инструмента, не требующего дополнительных затрат на инфраструктуру и '
        'администрирование. Внедрение клиентского веб-приложения, функционирующего '
        'непосредственно в браузере и хранящего данные локально, позволяет решить '
        'указанные проблемы при нулевых затратах на развёртывание.')

    add_body(doc,
        'Описанная проблема определяет актуальность настоящего проекта: разработка '
        'NexusCRM ориентирована именно на нишу лёгких клиентских CRM-инструментов, '
        'доступных без сторонней инфраструктуры.')

    # 1.2
    add_paragraph_heading(doc, '1.2', 'Анализ существующих CRM-решений')

    add_body(doc,
        'В целях обоснования подхода к разработке авторами проекта был проведён '
        'сравнительный анализ существующих CRM-систем по критериям, наиболее значимым '
        'для субъектов малого бизнеса. Ниже в таблице 1 представлены результаты анализа.')

    make_table1(doc)

    add_body(doc,
        'Анализ, приведённый в таблице 1, показывает, что ни одно из рассмотренных '
        'корпоративных CRM-приложений не обеспечивает возможности работы без '
        'интернет-соединения и без серверной инфраструктуры при полной функциональности. '
        'Bitrix24 и Salesforce характеризуются высокой сложностью освоения, а AmoCRM '
        'не предоставляет бесплатного плана. Разрабатываемый проект NexusCRM занимает '
        'нишу лёгкого клиентского инструмента с нулевой стоимостью и минимальными '
        'требованиями к окружению.')

    # 1.3
    add_paragraph_heading(doc, '1.3', 'Цель и задачи проекта')

    add_body(doc,
        'Целью проекта является разработка веб-приложения NexusCRM, обеспечивающего '
        'управление базой клиентов малого бизнеса с набором базовых CRM-функций: ведение '
        'записей, статусная модель, поиск, фильтрация и аналитика, без необходимости '
        'серверной части и постоянного интернет-соединения.')

    add_body(doc,
        'Для достижения цели авторами проекта были сформулированы следующие задачи:')
    for i, t in enumerate([
        'Провести анализ предметной области и требований к системе.',
        'Спроектировать многостраничную архитектуру клиентского приложения.',
        'Реализовать слой хранения данных на основе Web Storage API (localStorage).',
        'Разработать компонентную структуру JavaScript-логики с применением ES-модулей.',
        'Создать пользовательский интерфейс с тёмной цветовой схемой и адаптивной вёрсткой.',
        'Разработать модуль статистики с KPI-метриками и визуализацией данных.',
        'Провести тестирование и сформировать пользовательскую документацию.',
    ], 1):
        add_list_item(doc, t, numbered=i)

    add_pagebreak(doc)


def add_section2(doc):
    add_section(doc, '2', 'ПРОЕКТИРОВАНИЕ СИСТЕМЫ')

    # 2.1
    add_paragraph_heading(doc, '2.1', 'Концептуальная архитектура приложения')

    add_body(doc,
        'NexusCRM представляет собой клиентское многостраничное веб-приложение '
        '(Multi-Page Application, MPA), построенное на нативных веб-технологиях без '
        'применения клиентских фреймворков (React, Vue, Angular и подобных). '
        'Архитектура приложения реализована по принципу разделения ответственности: '
        'каждая HTML-страница отвечает за отдельную область функциональности, а '
        'JavaScript-файлы разделены по назначению и взаимодействуют посредством '
        'стандартного механизма ES-модулей (import/export).')

    add_body(doc,
        'На рисунке 1 представлена концептуальная архитектура приложения NexusCRM '
        'с разбивкой на три функциональных слоя.')

    if MPLOT:
        buf = make_architecture_diagram()
        embed_image(doc, buf, width_cm=13.5)
    add_figure_caption(doc, 1,
        'Концептуальная архитектура приложения NexusCRM',
        'составлено автором на основании проведённого исследования')

    add_body(doc,
        'Приложение состоит из трёх функциональных слоёв:')
    add_list_item(doc, 'слой представления \u2014 четыре HTML-страницы и единая '
                       'таблица стилей css/style.css;')
    add_list_item(doc, 'слой бизнес-логики \u2014 пять JavaScript ES-модулей, '
                       'реализующих прикладные операции и инфраструктурные функции;')
    add_list_item(doc, 'слой хранения данных \u2014 Web Storage API браузера '
                       '(localStorage), предоставляющий персистентное хранилище '
                       'без серверной части.')

    add_body(doc,
        'Взаимодействие слоёв организовано следующим образом: HTML-страницы '
        'подключают соответствующие JavaScript-модули через тег script с атрибутом '
        'type="module". Прикладные модули (main.js, form.js, stats.js) импортируют '
        'функции из инфраструктурных модулей storage.js и utils.js. Все операции '
        'записи и чтения данных клиентов инкапсулированы в storage.js и не '
        'дублируются в прикладной логике. Такое разделение обеспечивает '
        'поддерживаемость и расширяемость кода.')

    # 2.2
    add_paragraph_heading(doc, '2.2', 'Модель данных клиента')

    add_body(doc,
        'Каждый клиент в системе NexusCRM представлен JavaScript-объектом, '
        'хранящимся в виде элемента массива в localStorage браузера под ключом '
        '\u00abнexus_crm_clients\u00bb. В таблице 2 описана полная структура '
        'объекта клиента.')

    make_table2(doc)

    add_body(doc,
        'Уникальный идентификатор клиента формируется функцией generateId() из '
        'модуля utils.js как конкатенация текущей метки времени (Date.now()) и '
        'случайного четырёхзначного числа, что обеспечивает практическую уникальность '
        'в пределах одного браузера.')

    add_body(doc,
        'Статусная модель клиента предполагает три состояния:')
    add_list_item(doc, '\u00abНовый\u00bb (new) \u2014 клиент только добавлен в систему, '
                       'работа с ним не начата;')
    add_list_item(doc, '\u00abВ работе\u00bb (active) \u2014 ведётся активная работа '
                       'с клиентом;')
    add_list_item(doc, '\u00abЗавершён\u00bb (completed) \u2014 сделка закрыта, '
                       'работа завершена.')

    add_body(doc,
        'Переход между статусами осуществляется свободно, без ограничений на порядок: '
        'пользователь может в любой момент установить любой из трёх статусов через '
        'выпадающий список непосредственно в карточке клиента.')

    # 2.3
    add_paragraph_heading(doc, '2.3', 'Структура файлов проекта')

    add_body(doc,
        'Файловая структура проекта NexusCRM построена по принципу разделения '
        'статических ресурсов по типу: HTML-страницы расположены в корневой директории, '
        'таблицы стилей \u2014 в поддиректории css/, JavaScript-модули \u2014 '
        'в поддиректории js/. В таблице 3 представлены все файлы проекта '
        'с описанием их назначения.')

    make_table3(doc)

    add_body(doc,
        'Полная структура файлов проекта также приведена в Приложении 1.')

    add_pagebreak(doc)


def add_section3(doc):
    add_section(doc, '3', 'РЕАЛИЗАЦИЯ')

    # 3.1
    add_paragraph_heading(doc, '3.1', 'Стек технологий')

    add_body(doc,
        'Разработка приложения NexusCRM выполнена исключительно на нативных '
        'веб-технологиях без привлечения сторонних фреймворков, библиотек или '
        'систем сборки. Такой подход обеспечивает нулевое количество зависимостей, '
        'отсутствие необходимости в npm-пакетах и максимальную переносимость приложения. '
        'В таблице 4 представлен применяемый стек технологий.')

    make_table4(doc)

    add_body(doc,
        'Ключевой особенностью реализации является использование механизма ES-модулей '
        '(ECMAScript Modules), появившегося в стандарте ES2015 и поддерживаемого '
        'всеми современными браузерами без транспиляции. Это позволяет организовать '
        'код в отдельные файлы с явным описанием зависимостей через директивы import '
        'и export, аналогично модульным системам серверных платформ [4].')

    add_body(doc,
        'Данные приложения хранятся в localStorage \u2014 частном случае Web Storage '
        'API, предоставляемого браузером. localStorage обеспечивает персистентное '
        '(сохраняющееся между сессиями) хранение пар ключ\u2014значение с квотой '
        'обычно 5\u201310 МБ на домен. Данные сериализуются в формат JSON при '
        'сохранении и десериализуются при чтении [5].')

    # 3.2
    add_paragraph_heading(doc, '3.2', 'Описание программных модулей')

    add_body(doc,
        'Программный код приложения разделён на пять JavaScript-модулей в директории '
        'js/, каждый из которых выполняет строго определённые функции. '
        'Схема зависимостей модулей приведена в Приложении 2.')

    add_body(doc, 'Модуль storage.js реализует слой доступа к данным. '
                  'Он экспортирует следующие функции:', indent_first=True)
    add_list_item(doc, 'getClients() \u2014 чтение массива клиентов из localStorage '
                       'с валидацией структуры; при повреждении хранилища автоматически '
                       'восстанавливает демо-данные;')
    add_list_item(doc, 'saveClients(clients) \u2014 сохранение массива клиентов в '
                       'localStorage в формате JSON;')
    add_list_item(doc, 'initClients() \u2014 инициализация хранилища; при первом '
                       'запуске заполняет его шестью демонстрационными записями;')
    add_list_item(doc, 'addClient(client) \u2014 добавление нового объекта клиента '
                       'в конец массива;')
    add_list_item(doc, 'updateClient(id, updates) \u2014 частичное обновление полей '
                       'клиента по идентификатору (возвращает true при успехе);')
    add_list_item(doc, 'deleteClient(id) \u2014 удаление клиента по идентификатору;')
    add_list_item(doc, 'getClientById(id) \u2014 поиск и возврат клиента по '
                       'идентификатору или null;')
    add_list_item(doc, 'resetToDefaults() \u2014 служебная функция сброса '
                       'хранилища к демо-данным.')

    add_body(doc, 'Модуль utils.js содержит общие вспомогательные функции, '
                  'используемые несколькими модулями:')
    add_list_item(doc, 'generateId() \u2014 генерация уникального строкового '
                       'идентификатора на основе Date.now() и Math.random();')
    add_list_item(doc, 'formatDate(isoString) \u2014 преобразование ISO-даты '
                       'в формат ДД.ММ.ГГГГ;')
    add_list_item(doc, 'getInitials(fullName) \u2014 извлечение двух заглавных '
                       'инициалов из полного имени;')
    add_list_item(doc, 'getStatusLabel(status) / getStatusClass(status) \u2014 '
                       'получение русскоязычного обозначения и CSS-класса статуса;')
    add_list_item(doc, 'getAvatarColor(name) \u2014 детерминированный выбор цвета '
                       'аватара из палитры на основе хеш-функции от имени;')
    add_list_item(doc, 'escapeHtml(str) \u2014 экранирование HTML-спецсимволов '
                       'для защиты от XSS-атак [6];')
    add_list_item(doc, 'showToast(message, type, duration) \u2014 отображение '
                       'всплывающего уведомления с автоматическим закрытием;')
    add_list_item(doc, 'debounce(fn, delay) \u2014 создание отложенной версии функции '
                       'для снижения частоты вызовов;')
    add_list_item(doc, 'setActiveNav() \u2014 подсветка активного пункта навигации '
                       'по текущему URL.')

    add_body(doc,
        'Модуль main.js содержит логику главной страницы (index.html). '
        'Он реализует рендеринг карточек клиентов функцией buildClientCard(), '
        'применение фильтров и поиска функцией filterClients(), обработку изменения '
        'статуса (handleStatusChange()) и удаления клиентов (handleDelete()). '
        'Поиск реализован с задержкой 280 мс посредством функции debounce() '
        'для снижения нагрузки при наборе текста. Модуль поддерживает восстановление '
        'страницы из кэша браузера (bfcache): при событии pageshow с persisted=true '
        'вызывается облегчённая функция reinit(), которая обновляет данные без '
        'повторной привязки обработчиков событий.')

    add_body(doc,
        'Модуль form.js отвечает за логику страницы form.html. '
        'Он поддерживает два режима: создание нового клиента и редактирование '
        'существующего. Режим редактирования определяется по параметру URL '
        '\u00abedit=<id>\u00bb: если параметр присутствует, форма переходит '
        'в режим редактирования и заполняется текущими данными клиента через '
        'getClientById(). Реализована клиентская валидация: проверка заполненности '
        'обязательных полей (fullName, phone, email) и корректности формата '
        'электронной почты регулярным выражением.')

    add_body(doc,
        'Модуль stats.js формирует данные для страницы статистики (stats.html). '
        'Функция renderKPICards() рассчитывает пять показателей: общее количество '
        'клиентов, количество новых, активных, завершённых и показатель конверсии '
        '(отношение завершённых к общему числу в процентах). '
        'Функция renderStatusBreakdown() формирует разбивку по статусам с '
        'анимированными прогресс-барами. Функция renderMonthlyActivity() строит '
        'столбчатую диаграмму активности за последние шесть месяцев на основе '
        'поля createdAt клиентов.')

    # 3.3
    add_paragraph_heading(doc, '3.3', 'Пользовательский интерфейс')

    add_body(doc,
        'Приложение NexusCRM включает четыре страницы с единой навигацией и '
        'общей таблицей стилей. Интерфейс выполнен в тёмной цветовой схеме '
        '(dark mode) с акцентным фиолетово-синим цветом (#6366f1) и '
        'вспомогательным пурпурным (#8b5cf6).')

    add_body(doc,
        'Страница \u00abКлиенты\u00bb (index.html) является главной страницей '
        'приложения. Она содержит секцию-hero с заголовком и счётчиками записей, '
        'панель инструментов (toolbar) с полем поиска и фильтрами по статусу, '
        'а также сетку карточек клиентов. Каждая карточка отображает: '
        'аватар с инициалами (цвет детерминирован именем), полное имя, '
        'адрес электронной почты в виде кликабельной ссылки (mailto:), '
        'статус-значок (badge), телефон, текстовый комментарий и дату добавления. '
        'В подвале карточки расположены элементы управления: выпадающий список '
        'смены статуса, кнопка редактирования и кнопка удаления.')

    add_body(doc,
        'Страница \u00abДобавить клиента\u00bb (form.html) реализует форму '
        'с полями: полное имя, телефон, email (все три обязательны), '
        'статус (выпадающий список с тремя значениями), дата добавления и '
        'текстовый комментарий (до 500 символов). Форма поддерживает два режима: '
        'создание нового клиента и редактирование существующего. '
        'При ошибке валидации поле помечается CSS-классом field-error '
        'и отображается уведомление-тост с текстом ошибки.')

    add_body(doc,
        'Страница \u00abСтатистика\u00bb (stats.html) отображает аналитические '
        'данные по базе клиентов: пять KPI-карточек (всего клиентов, новых, '
        'в работе, завершённых, показатель конверсии), разбивку по статусам '
        'с анимированными прогресс-барами и столбчатую диаграмму активности '
        'по последним шести месяцам.')

    add_body(doc,
        'Страница \u00abО проекте\u00bb (about.html) содержит описание шести '
        'ключевых возможностей приложения в виде карточек, пошаговую инструкцию '
        'по использованию и блок с информацией об авторе.')

    add_body(doc,
        'На рисунке 2 представлена схема навигационных переходов между '
        'страницами приложения.')

    if MPLOT:
        buf = make_navigation_diagram()
        embed_image(doc, buf, width_cm=13.5)
    add_figure_caption(doc, 2,
        'Схема переходов между страницами приложения NexusCRM',
        'составлено автором на основании проведённого исследования')

    add_body(doc,
        'Все страницы адаптированы для мобильных устройств: при ширине экрана '
        'менее 768 пикселей навигационное меню преобразуется в выпадающее, '
        'активируемое кнопкой-гамбургером. Сетка карточек клиентов '
        'перестраивается с трёх столбцов до одного при уменьшении экрана.')

    # 3.4
    add_paragraph_heading(doc, '3.4', 'Алгоритм работы приложения')

    add_body(doc,
        'На рисунке 3 представлен алгоритм инициализации и работы '
        'главной страницы приложения.')

    if MPLOT:
        buf = make_algorithm_diagram()
        embed_image(doc, buf, width_cm=9.0)
    add_figure_caption(doc, 3,
        'Алгоритм инициализации главной страницы приложения',
        'составлено автором на основании проведённого исследования')

    add_body(doc,
        'Работа приложения начинается с загрузки HTML-страницы в браузере. '
        'Поскольку JS-модули подключаются с атрибутом type="module" и '
        'являются отложенными (deferred), к моменту выполнения скрипта DOM '
        'уже полностью разобран. Функция init() последовательно выполняет '
        'следующие действия:')
    for i, t in enumerate([
        'вызывает initClients() для инициализации хранилища данных '
        '(при пустом хранилище — автоматическое заполнение демо-данными);',
        'вызывает setActiveNav() для подсветки активного пункта навигации;',
        'вызывает renderClients() для отображения карточек клиентов '
        'с учётом текущего фильтра и строки поиска;',
        'выполняет однократную привязку обработчиков событий (поиск, '
        'фильтры, мобильное меню) через флаг listenersAttached.',
    ], 1):
        add_list_item(doc, t, numbered=i)

    add_body(doc,
        'При восстановлении страницы из кэша браузера (событие pageshow '
        'с persisted=true) вызывается облегчённая функция reinit(), '
        'которая обновляет отображение без повторной привязки обработчиков '
        'во избежание их дублирования.')

    add_body(doc,
        'Функция filterClients() применяет текущий фильтр статуса '
        'и строку поиска к массиву клиентов. Поиск выполняется '
        'по полям fullName, email и phone с применением метода '
        'String.prototype.includes() после приведения к нижнему регистру '
        'посредством toLowerCase(). Оба условия (фильтр статуса и поиск) '
        'применяются одновременно.')

    add_pagebreak(doc)


def add_section4(doc):
    add_section(doc, '4', 'РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ')

    # 4.1
    add_paragraph_heading(doc, '4.1', 'Требования к запуску и установке')

    add_body(doc,
        'NexusCRM не требует установки программного обеспечения, настройки '
        'серверного окружения или постоянного подключения к интернету. '
        'Данные хранятся локально в браузере.')

    add_body(doc,
        'Минимальные системные требования:')
    add_list_item(doc, 'современный веб-браузер с поддержкой ES-модулей '
                       '(type="module") и Web Storage API;')
    add_list_item(doc, 'операционная система: Windows 10+, macOS 10.14+, '
                       'Linux (любая актуальная версия дистрибутива);')
    add_list_item(doc, 'минимальная версия браузера: Google Chrome 80+, '
                       'Mozilla Firefox 80+, Microsoft Edge 80+, Safari 14+.')

    add_body(doc,
        'Порядок запуска приложения:')
    for i, t in enumerate([
        'Скачать или клонировать репозиторий проекта на локальный диск.',
        'Открыть файл index.html в поддерживаемом веб-браузере.',
        'При первом запуске хранилище автоматически заполнится '
        'шестью демонстрационными записями клиентов.',
    ], 1):
        add_list_item(doc, t, numbered=i)

    add_body(doc,
        'В некоторых конфигурациях браузера при открытии файла '
        'по протоколу file:// ES-модули могут быть заблокированы '
        'политикой CORS. В этом случае рекомендуется запустить '
        'локальный HTTP-сервер следующей командой в директории проекта: '
        'python -m http.server 8000, после чего открыть адрес '
        'http://localhost:8000 в браузере.')

    # 4.2
    add_paragraph_heading(doc, '4.2', 'Основные сценарии использования')

    add_body(doc,
        'Ниже описаны основные сценарии работы с приложением NexusCRM.')

    add_body(doc,
        'Сценарий 1. Просмотр базы клиентов. '
        'Пользователь открывает страницу \u00abКлиенты\u00bb (index.html). '
        'При первом запуске хранилище автоматически заполняется шестью '
        'демонстрационными записями. Отображается сетка карточек '
        'с общим счётчиком и счётчиком видимых записей.')

    add_body(doc,
        'Сценарий 2. Добавление нового клиента. '
        'Пользователь нажимает кнопку \u00abДобавить клиента\u00bb. '
        'Открывается страница form.html в режиме создания. '
        'Пользователь заполняет форму (имя, телефон и email обязательны), '
        'выбирает статус, при необходимости вводит комментарий и дату, '
        'нажимает \u00abСохранить клиента\u00bb. '
        'После сохранения отображается уведомление-тост и выполняется '
        'перенаправление на главную страницу.')

    add_body(doc,
        'Сценарий 3. Поиск и фильтрация клиентов. '
        'Пользователь вводит текст в поле поиска. '
        'Поиск осуществляется по имени, email и телефону '
        'с задержкой 280 мс (debounce). '
        'Одновременно можно выбрать один из фильтров статуса: '
        '\u00abВсе\u00bb, \u00abНовые\u00bb, \u00abВ работе\u00bb, '
        '\u00abЗавершённые\u00bb. Оба условия применяются одновременно. '
        'Кнопка очистки поиска (иконка \u00d7) сбрасывает строку запроса.')

    add_body(doc,
        'Сценарий 4. Изменение статуса клиента. '
        'В подвале карточки клиента пользователь выбирает новый статус '
        'из выпадающего списка. Изменение немедленно сохраняется '
        'в localStorage. Значок статуса в карточке обновляется '
        'без полной перерисовки сетки. '
        'Отображается уведомление-тост с новым статусом.')

    add_body(doc,
        'Сценарий 5. Редактирование данных клиента. '
        'Пользователь нажимает кнопку редактирования '
        '(иконка карандаша) в карточке клиента. '
        'Открывается страница form.html с параметром ?edit=<id>. '
        'Форма автоматически заполняется текущими данными клиента. '
        'После внесения изменений и нажатия \u00abСохранить изменения\u00bb '
        'данные обновляются в localStorage и выполняется перенаправление '
        'на главную страницу.')

    add_body(doc,
        'Сценарий 6. Удаление клиента. '
        'Пользователь нажимает кнопку удаления (иконка корзины) '
        'в карточке клиента. Отображается стандартный диалог подтверждения '
        'браузера. После подтверждения карточка удаляется с анимацией '
        'исчезновения, данные удаляются из localStorage, '
        'отображается уведомление-тост.')

    add_body(doc,
        'Сценарий 7. Просмотр статистики. '
        'Пользователь переходит на страницу \u00abСтатистика\u00bb. '
        'Отображаются: пять KPI-карточек (общее количество, новые, '
        'в работе, завершённые, показатель конверсии), разбивка '
        'по статусам с анимированными прогресс-барами и '
        'столбчатая гистограмма активности за последние шесть месяцев.')

    # 4.3
    add_paragraph_heading(doc, '4.3', 'Ограничения и перспективы развития')

    add_body(doc,
        'К текущим ограничениям системы относятся следующие:')
    add_list_item(doc, 'данные хранятся только в localStorage конкретного '
                       'браузера и не синхронизируются между устройствами;')
    add_list_item(doc, 'объём хранимых данных ограничен квотой localStorage '
                       'браузера (как правило, 5\u201310 МБ);')
    add_list_item(doc, 'отсутствует разграничение прав доступа пользователей;')
    add_list_item(doc, 'система не поддерживает экспорт данных в табличные '
                       'форматы (CSV, XLSX);')
    add_list_item(doc, 'аналитический модуль не поддерживает произвольный '
                       'выбор периода наблюдения.')

    add_body(doc,
        'В качестве перспектив развития проекта авторами проекта '
        'предлагаются следующие направления:')
    add_list_item(doc, 'реализация синхронизации данных через облачное '
                       'хранилище (Firebase Realtime Database или Supabase);')
    add_list_item(doc, 'добавление функциональности экспорта клиентской базы '
                       'в форматы CSV и XLSX;')
    add_list_item(doc, 'разработка серверной части на Node.js или Python '
                       'для многопользовательского режима;')
    add_list_item(doc, 'реализация режима Progressive Web App (PWA) '
                       'для установки приложения на устройство;')
    add_list_item(doc, 'добавление системы тегов, расширенных полей профиля '
                       'клиента и напоминаний.')

    add_pagebreak(doc)


def add_conclusion(doc):
    add_structural(doc, 'ЗАКЛЮЧЕНИЕ')

    add_body(doc,
        'В ходе настоящей проектной работы было разработано клиентское '
        'веб-приложение NexusCRM \u2014 система управления базой клиентов '
        'для малого бизнеса, реализованная на нативных технологиях HTML5, '
        'CSS3 и JavaScript с применением ES-модулей.')

    add_body(doc,
        'В процессе работы были решены все поставленные задачи: '
        'проведён анализ предметной области и существующих CRM-решений, '
        'спроектирована и реализована многостраничная архитектура приложения, '
        'разработан слой хранения данных на основе Web Storage API, '
        'создан адаптивный пользовательский интерфейс с тёмной цветовой схемой, '
        'реализованы все операции управления клиентскими данными, '
        'а также модуль статистики и аналитики.')

    add_body(doc,
        'Разработанное приложение обладает следующими достоинствами: '
        'полная независимость от серверной инфраструктуры и интернет-соединения '
        'при работе с данными, отсутствие зависимостей от сторонних библиотек и '
        'систем сборки, поддержка мобильных устройств, '
        'защита от XSS-атак посредством экранирования HTML-спецсимволов, '
        'автоматическое восстановление демо-данных при повреждении хранилища, '
        'а также поддержка кэширования страниц браузером (bfcache).')

    add_body(doc,
        'Практическая значимость проекта определяется его готовностью '
        'к применению в реальных условиях малого бизнеса: '
        'приложение может быть размещено на любом статическом хостинге '
        'без дополнительных затрат на инфраструктуру.')

    add_body(doc,
        'Перспективы развития проекта связаны с добавлением серверной части '
        'для многопользовательского режима, облачной синхронизацией данных, '
        'функциональностью экспорта и расширенными аналитическими инструментами.')

    add_pagebreak(doc)


def add_references(doc):
    add_structural(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')

    refs = [
        'Buttle F., Maklan S. (2019). Customer Relationship Management: '
        'Concepts and Technologies. Routledge. 448 p.',

        'MDN Web Docs. (2024). Using the Web Storage API. Mozilla. '
        'https://developer.mozilla.org/en-US/docs/Web/API/Web_Storage_API/Using_the_Web_Storage_API',

        'Salesforce Research. (2022). State of Sales Report (5th ed.). Salesforce. '
        'https://www.salesforce.com/resources/research-reports/state-of-sales/',

        'MDN Web Docs. (2024). JavaScript modules. Mozilla. '
        'https://developer.mozilla.org/en-US/docs/Web/JavaScript/Guide/Modules',

        'MDN Web Docs. (2024). Window.localStorage. Mozilla. '
        'https://developer.mozilla.org/en-US/docs/Web/API/Window/localStorage',

        'OWASP. (2024). Cross Site Scripting (XSS). OWASP Foundation. '
        'https://owasp.org/www-community/attacks/xss/',

        'Simpson K. (2020). You Don\u2019t Know JS Yet: Get Started. '
        'Independently published. 143 p.',

        'Frain B. (2022). Responsive Web Design with HTML5 and CSS. '
        'Packt Publishing. 476 p.',

        'MDN Web Docs. (2024). Page lifecycle API: bfcache. Mozilla. '
        'https://developer.mozilla.org/en-US/docs/Web/Performance/Guides/bfcache',

        'Google Developers. (2024). Google Fonts API v1. Google. '
        'https://developers.google.com/fonts/docs/getting_started',
    ]

    for i, ref in enumerate(refs, 1):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True)
        r = p.add_run(f'{i} {ref}')
        _set_run(r)

    add_pagebreak(doc)


def add_appendices(doc):
    # Приложение 1
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, indent_first=False)
    r = p.add_run('Приложение 1')
    _set_run(r)

    add_blank(doc)
    p2 = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r2 = p2.add_run('Полная файловая структура проекта NexusCRM')
    _set_run(r2, bold=True)
    add_blank(doc)

    tree = (
        'meneger/\n'
        '\u251c\u2500\u2500 about.html\n'
        '\u251c\u2500\u2500 css/\n'
        '\u2502   \u2514\u2500\u2500 style.css\n'
        '\u251c\u2500\u2500 docs/\n'
        '\u2502   \u251c\u2500\u2500 generate_doc.py\n'
        '\u2502   \u251c\u2500\u2500 NexusCRM_Documentation.docx\n'
        '\u2502   \u2514\u2500\u2500 NexusCRM_Documentation.pdf\n'
        '\u251c\u2500\u2500 form.html\n'
        '\u251c\u2500\u2500 index.html\n'
        '\u251c\u2500\u2500 js/\n'
        '\u2502   \u251c\u2500\u2500 form.js\n'
        '\u2502   \u251c\u2500\u2500 main.js\n'
        '\u2502   \u251c\u2500\u2500 stats.js\n'
        '\u2502   \u251c\u2500\u2500 storage.js\n'
        '\u2502   \u2514\u2500\u2500 utils.js\n'
        '\u251c\u2500\u2500 README.md\n'
        '\u2514\u2500\u2500 stats.html'
    )

    for line in tree.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before     = Pt(0)
        p.paragraph_format.space_after      = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(10)
        _rpr_font(r._r, 'Courier New')

    add_pagebreak(doc)

    # Приложение 2
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, indent_first=False)
    r = p.add_run('Приложение 2')
    _set_run(r)

    add_blank(doc)
    p2 = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r2 = p2.add_run('Схема зависимостей JavaScript-модулей')
    _set_run(r2, bold=True)
    add_blank(doc)

    if MPLOT:
        buf = make_modules_diagram()
        embed_image(doc, buf, width_cm=13.0)
        p_note = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
        r_note = p_note.add_run(
            'Примечание \u2013 составлено автором на основании проведённого исследования')
        _set_run(r_note)
    else:
        add_body(doc,
            'Схема зависимостей не сформирована: требуется установка matplotlib.')


def make_modules_diagram():
    """Приложение 2 — Схема зависимостей модулей."""
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_facecolor('white')

    def box(x, y, w, h, color, label):
        r = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.15',
                           linewidth=1.3, edgecolor='#444', facecolor=color)
        ax.add_patch(r)
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                fontsize=9, fontweight='bold')

    def arr(x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#333', lw=1.4))

    # Инфраструктура
    box(0.5, 0.5, 2.0, 0.8, '#dbeafe', 'storage.js')
    box(3.0, 0.5, 2.0, 0.8, '#dbeafe', 'utils.js')

    # Прикладные
    box(0.0, 2.5, 2.0, 0.8, '#e0e7ff', 'main.js')
    box(2.5, 2.5, 2.0, 0.8, '#e0e7ff', 'form.js')
    box(5.0, 2.5, 2.0, 0.8, '#e0e7ff', 'stats.js')

    # HTML страницы
    box(0.0, 4.0, 2.0, 0.7, '#d1fae5', 'index.html')
    box(2.5, 4.0, 2.0, 0.7, '#d1fae5', 'form.html')
    box(5.0, 4.0, 2.0, 0.7, '#d1fae5', 'stats.html')
    box(7.5, 4.0, 2.0, 0.7, '#d1fae5', 'about.html')

    # storage.js → all app modules
    arr(1.5, 1.3, 1.0, 2.5)
    arr(1.5, 1.3, 3.5, 2.5)
    arr(1.5, 1.3, 6.0, 2.5)

    # utils.js → all app modules
    arr(4.0, 1.3, 1.0, 2.5)
    arr(4.0, 1.3, 3.5, 2.5)
    arr(4.0, 1.3, 6.0, 2.5)

    # app modules → HTML
    arr(1.0, 3.3, 1.0, 4.0)
    arr(3.5, 3.3, 3.5, 4.0)
    arr(6.0, 3.3, 6.0, 4.0)

    # Легенда
    legend_items = [
        mpatches.Patch(facecolor='#d1fae5', edgecolor='#444', label='HTML-страницы'),
        mpatches.Patch(facecolor='#e0e7ff', edgecolor='#444', label='Прикладные модули'),
        mpatches.Patch(facecolor='#dbeafe', edgecolor='#444', label='Инфраструктурные модули'),
    ]
    ax.legend(handles=legend_items, loc='center right', fontsize=8,
              framealpha=0.9, bbox_to_anchor=(11.0, 2.5))

    ax.set_title('Схема зависимостей JavaScript-модулей NexusCRM',
                 fontsize=11, pad=10)
    return fig_to_buf(fig)


# ══════════════════════════════════════════════════════════════════════════════
# ГЛАВНАЯ ФУНКЦИЯ
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    print('[1/4] Создание документа...')
    doc = setup_document()

    print('[2/4] Формирование содержимого...')
    add_title_page(doc)
    add_contents_page(doc)
    add_introduction(doc)
    add_section1(doc)
    add_section2(doc)
    add_section3(doc)
    add_section4(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendices(doc)

    # Нумерация страниц
    add_page_numbers(doc)
    hide_first_page_number(doc)

    print(f'[3/4] Saving .docx -> {OUT_DOCX}')
    doc.save(str(OUT_DOCX))
    print(f'      OK: {OUT_DOCX}')
    return doc


def convert_to_pdf():
    print(f'[4/4] Converting to PDF -> {OUT_PDF}')
    try:
        from docx2pdf import convert
        convert(str(OUT_DOCX), str(OUT_PDF))
        print(f'      OK: {OUT_PDF}')
        return True
    except Exception as e:
        print(f'      [WARN] docx2pdf: {e}')

    # Запасной вариант — LibreOffice
    try:
        import subprocess
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf',
             '--outdir', str(DOCS_DIR), str(OUT_DOCX)],
            capture_output=True, timeout=60
        )
        if result.returncode == 0:
            print(f'      OK (LibreOffice): {OUT_PDF}')
            return True
        print(f'      [WARN] LibreOffice failed: {result.stderr.decode()}')
    except Exception as e:
        print(f'      [WARN] LibreOffice: {e}')

    print('\n  PDF не создан автоматически.')
    print('  Откройте NexusCRM_Documentation.docx в Microsoft Word')
    print('  и сохраните как PDF через меню Файл → Сохранить как.')
    return False


if __name__ == '__main__':
    build_document()
    convert_to_pdf()
    print('\nГотово!')
